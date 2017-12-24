from docx import Document
import nltk
import numpy as np
import hashlib
import importlib as im
import matplotlib.pyplot as plt
l=['nltk','docx','numpy','matplotlib']#required modules 
for x in l:
    try:
        im.import_module(x) #check for modules
    except ImportError:
        import pip
        print('Installing... ' + x) #if not found install modules
        pip.main(['install',x])
        globals()[x]=im.import_module(x)
        print('Install Successful')
try:
    do1=Document('hosp.docx') #provide file name 1
except:
    print('No such file')
try:
    do2=Document('foo.docx') #provide file name 2
except:
    print('No such file')
fullText1=[] #for storing text as a list
fullText2=[]
for para in do1.paragraphs:
        fullText1.append(para.text) #appending new text to list
for para in do2.paragraphs:
        fullText2.append(para.text)
str1=' '.join(fullText1) #convert list to string for tokenizing
str2=' '.join(fullText2) 
token1 = nltk.word_tokenize(str1) #tokenize for analysis
token2 = nltk.word_tokenize(str2)
arr1=np.zeros(shape=(1000000,1),dtype=int) #np array for hashing
arr2=np.zeros(shape=(1000000,1),dtype=int)
arr=np.zeros(shape=(1000000,1),dtype=int)
for i in range(0,len(token1)):
    hval=int(hashlib.sha256(token1[i].encode('utf-8')).hexdigest(), 16) % 10**5 #compute hash value
    arr1[hval][0]=arr1[hval][0]+1
    arr[hval][0]=arr[hval][0]+1 #store hash value
for i in range(0,len(token2)):
    hval=int(hashlib.sha256(token2[i].encode('utf-8')).hexdigest(), 16) % 10**5
    arr2[hval][0]=arr2[hval][0]+1
    arr[hval][0]=arr[hval][0]+1
d=0
hit=0
for i in range(0,99999):
    if arr[i][0]!=0: #total number of distinct words
        d=d+1
    if arr1[i][0]==arr2[i][0] and arr1[i][0]!=0 and arr2[i][0]!=0: #number of matches
        hit=hit+1
p=(float(hit)/float(d))*100.0
print('HIT PERCENTAGE ...%.2f '  %p)
plt.plot(arr1,arr2)
plt.ylabel('Hash Index') #graph plotting
plt.show()  #NUMBER OF LINES ORIGINATING IS INVERSLY PROPORTIONAL TO MATCH PERCENTAGE 
